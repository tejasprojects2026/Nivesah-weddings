from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "output/doc/Nivesah_Weddings_Digital_Implementation_Report.docx"
INK = "14110D"
CHAMPAGNE = "D9C29C"
OLIVE = "6B6B45"
WARM = "F7F3EB"
MUTED = "6C6861"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
for name, size, color in [("Title", 31, INK), ("Heading 1", 19, INK), ("Heading 2", 13, OLIVE)]:
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def title(text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(INK)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.paragraph_format.space_after = Pt(18)
        p2.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

def heading(text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(7)

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], INK)
        cell_text(t.rows[0].cells[i], h, bold=True, color="FFFFFF")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if len(t.rows) % 2 == 0:
                shade(cells[i], WARM)
            cell_text(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nivesah Weddings - Digital Implementation Report | 14 July 2026")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)

footer(section)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(42)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NIVESAH WEDDINGS")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string(CHAMPAGNE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Digital Implementation\nReport")
r.bold = True
r.font.name = "Aptos Display"
r.font.size = Pt(33)
r.font.color.rgb = RGBColor.from_string(INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SEO, Google Analytics and Enquiry Form Integration")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MUTED)
doc.add_paragraph().paragraph_format.space_after = Pt(42)
table(["Prepared for", "Report date", "Website"], [["Nivesah Weddings", "14 July 2026", "www.nivesahweddings.in"]], [2.15, 1.75, 2.65])
heading("Executive summary")
body("The website has been configured with foundational search-engine optimisation, Google Analytics tracking, and a working online enquiry workflow. Google Search Console has successfully accepted the sitemap and identified eight URLs for crawling.")
table(["Area", "Status", "Outcome"], [
    ["SEO and sitemap", "Completed", "Sitemap submitted successfully; 8 URLs discovered."],
    ["Google Analytics", "Implemented", "Google tag added using Measurement ID G-25WMCG5LM7."],
    ["Enquiry form", "Implemented", "Form sends enquiries through Web3Forms with user feedback."],
], [1.6, 1.2, 3.75])

doc.add_page_break()
title("1. Search Engine Optimisation", "Technical SEO configuration and Search Console status")
heading("Completed configuration")
bullets([
    "Search-friendly page title and meta description for Nivesah Weddings and iFilms Media services.",
    "Robots directives allow search engines to crawl and index the public site.",
    "XML sitemap added at /sitemap.xml and referenced in /robots.txt.",
    "Canonical, Open Graph and Twitter metadata are present to support preferred URLs and rich sharing previews.",
    "Local business context is included in page metadata for Pune, Maharashtra, India.",
])
heading("Google Search Console")
table(["Item", "Current result"], [
    ["Submitted sitemap", "https://www.nivesahweddings.in/sitemap.xml"],
    ["Sitemap status", "Success"],
    ["Discovered pages", "8"],
    ["Google page-indexing report", "Processing - expected for a newly submitted sitemap"],
], [2.2, 4.35])
body("What this means: Google can access the sitemap and has queued the listed URLs for crawling. Discovery does not guarantee immediate search ranking or indexing; Google controls crawl and indexing timing.")
heading("Recommended next steps")
bullets([
    "Allow 24-48 hours for Search Console's Page Indexing report to populate, then review any exclusions or errors.",
    "Use URL Inspection to request indexing for the homepage and any future important standalone pages.",
    "Keep the preferred www domain consistent across canonical tags, sitemap URLs, redirects and social links.",
    "Publish dedicated, indexable pages for major services or portfolio stories when ready. Hash links such as #about are sections of the homepage, not separate pages in Google.",
])

doc.add_page_break()
title("2. Google Analytics", "Visitor measurement implementation")
heading("Implementation details")
table(["Setting", "Configured value"], [
    ["Platform", "Google Analytics 4 (GA4)"],
    ["Measurement ID", "G-25WMCG5LM7"],
    ["Tag placement", "Global document head - loads on all website pages"],
    ["Deployment", "Committed and pushed to the main GitHub branch"],
], [2.1, 4.45])
body("The Google tag is loaded asynchronously, so it does not intentionally block page rendering. It initializes the GA4 configuration when a visitor loads the site.")
heading("How to verify after deployment")
bullets([
    "Open the live website in a normal browser window after the host has deployed the latest main-branch commit.",
    "In Google Analytics, open Reports > Realtime. The visit should appear within a few minutes.",
    "Use a separate device or mobile data if your own traffic is filtered in the Analytics property.",
    "Do not use browser ad blockers during the test; they can block analytics requests.",
])
heading("Data available once traffic arrives")
body("GA4 can report visits, traffic source, device type, geography, engagement and pages viewed. The current setup provides baseline page-view measurement. Event tracking for actions such as enquiry submission, WhatsApp clicks, phone taps and Instagram clicks can be added as a next phase.")

doc.add_page_break()
title("3. Enquiry Form Integration", "Lead capture through Web3Forms")
heading("Form workflow")
body("The website's 'Tell us about your event' form is connected to Web3Forms using the supplied access key. Visitors submit the form directly from the website; the form no longer redirects them to WhatsApp after submission.")
table(["Field captured", "Included in enquiry"], [
    ["Name", "Yes"],
    ["Phone number", "Yes"],
    ["Event date", "Yes"],
    ["Event location", "Yes"],
    ["Service required", "Yes"],
    ["Message", "Yes"],
], [2.55, 4.0])
heading("Customer experience")
bullets([
    "Required-field validation runs before the form is sent.",
    "The button changes to 'Sending...' while Web3Forms processes the request.",
    "A clear success message confirms a sent enquiry and resets the form.",
    "A clear error message asks the visitor to retry or use WhatsApp if delivery fails.",
])
heading("Client verification checklist")
bullets([
    "After deployment, submit one real test enquiry using a monitored email address and phone number.",
    "Confirm the enquiry arrives at the recipient email configured in the Web3Forms dashboard for this access key.",
    "Check the email subject: 'New Nivesah Weddings enquiry'.",
    "If a test does not arrive, confirm the Web3Forms access key's recipient email and domain settings in the Web3Forms dashboard.",
])
heading("Delivery note")
body("The access key is designed to be used in a browser-based form. Receipt and routing of submissions are controlled by the email and settings attached to the key in Web3Forms. A live test after deployment is required to confirm final email delivery.")

doc.add_page_break()
title("4. Handover and Next Steps", "Items for the client team")
heading("Current project delivery")
table(["Component", "Delivery state", "Client action"], [
    ["Website source changes", "Committed and pushed to GitHub main", "Confirm hosting deployment completes."],
    ["Sitemap", "Accepted by Google Search Console", "Review Page Indexing after 24-48 hours."],
    ["Analytics", "Tag added", "Check GA4 Realtime after deployment."],
    ["Web3Forms", "Form integration added", "Send and confirm one live test enquiry."],
], [1.6, 2.15, 2.8])
heading("Recommended phase two")
bullets([
    "Configure GA4 conversion events for successful enquiry submissions and key contact actions.",
    "Add separate service and portfolio-story URLs with unique content, title tags and descriptions.",
    "Add structured data for the local business, services, reviews and images where appropriate.",
    "Review Search Console monthly for index coverage, search queries, clicks and Core Web Vitals.",
    "Create a simple monthly report covering organic traffic, enquiries, top search terms and next SEO opportunities.",
])
heading("Important expectation")
body("SEO creates the technical foundation for Google to find and understand the website. Meaningful organic visibility usually develops over time as Google crawls the site and the site earns relevance, content depth and links. Analytics and form testing should be validated on the published site after the hosting provider completes deployment.")

doc.save(OUT)
print(OUT)
